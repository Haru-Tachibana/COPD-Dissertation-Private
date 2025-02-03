import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Function to add a table to the document
def add_table_to_doc(df, doc):
    # Add a table
    table = doc.add_table(rows=1, cols=len(df.columns))

    # Add headers and set them to bold and centered
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = col_name
        # Set header text to bold
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        # Center the header text
        hdr_cells[i].paragraphs[0].alignment = 1  # 1 is center alignment

    # Set the table border style using the XML (with correct namespace handling)
    tbl = table._element
    tblBorders = tbl.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    
    if tblBorders is not None:
        for border in tblBorders:
            if border.tag.endswith('top') or border.tag.endswith('bottom'):
                border.set(qn('w:val'), 'single')
                border.set(qn('w:space'), '0')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
    else:
        # If tblBorders is not found, we create it
        tblBorders = OxmlElement('w:tblBorders')
        tbl.append(tblBorders)
        top_border = OxmlElement('w:top')
        tblBorders.append(top_border)
        top_border.set(qn('w:val'), 'single')
        top_border.set(qn('w:space'), '0')
        top_border.set(qn('w:sz'), '4')
        top_border.set(qn('w:space'), '0')

        bottom_border = OxmlElement('w:bottom')
        tblBorders.append(bottom_border)
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:space'), '0')
        bottom_border.set(qn('w:sz'), '4')
        bottom_border.set(qn('w:space'), '0')

    # Add rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

# Load the existing Word document
doc_path = "/Users/yangyangxiayule/Documents/GitHub/COPD-Project/TABLES.docx"
doc = Document(doc_path)

# Example: Load CSV and add new table
df = pd.read_csv("/Users/yangyangxiayule/Documents/GitHub/COPD-Project/bottom_copd_counties.csv")
add_table_to_doc(df, doc)

# If you want to add another table, load another CSV and add it:
# df2 = pd.read_csv("another_csv_file.csv")
# add_table_to_doc(df2, doc)

# Save the updated document
doc.save(doc_path)
